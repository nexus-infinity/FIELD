#!/usr/bin/env python3
"""
BERJAK | NEXUS | INFINITY FRE Document Processor
Handles extraction of text content from various file types
"""

import re
import PyPDF2
from pathlib import Path
from typing import Dict, Optional, List, Tuple
import csv
import docx
import openpyxl
import xlrd
import pandas as pd
import pyodbc
import cv2
import numpy as np
from PIL import Image
import pytesseract
import face_recognition
from deepface import DeepFace
import exifread
import magic
from datetime import datetime
import json
import hashlib

class DocumentProcessor:
    def __init__(self):
        self.processed_files = {}
        self.extraction_stats = {
            'pdf': {'success': 0, 'failed': 0},
            'doc': {'success': 0, 'failed': 0},
            'xlsx': {'success': 0, 'failed': 0},
            'xls': {'success': 0, 'failed': 0},
            'csv': {'success': 0, 'failed': 0},
            'txt': {'success': 0, 'failed': 0},
            'myo': {'success': 0, 'failed': 0},
            'img': {'success': 0, 'failed': 0}
        }

    def process_file(self, file_path: Path) -> Dict:
        """Process a single file and extract relevant information"""
        try:
            content = self._extract_content(file_path)
            metadata = self._extract_metadata(file_path)
            entities = self._extract_entities(content)
            dates = self._extract_dates(content)
            
            result = {
                'content': content,
                'metadata': metadata,
                'entities': entities,
                'dates': dates,
                'processed_at': datetime.now().isoformat()
            }
            
            self.processed_files[str(file_path)] = result
            return result
            
        except Exception as e:
            print(f"Error processing {file_path}: {str(e)}")
            return None

    def _extract_content(self, file_path: Path) -> str:
        """Extract text content from various file types"""
        suffix = file_path.suffix.lower()
        
        if suffix == '.pdf':
            return self._extract_pdf_content(file_path)
        elif suffix in ['.doc', '.docx']:
            return self._extract_doc_content(file_path)
        elif suffix == '.xlsx':
            return self._extract_excel_content(file_path)
        elif suffix == '.xls':
            return self._extract_legacy_excel_content(file_path)
        elif suffix == '.csv':
            return self._extract_csv_content(file_path)
        elif suffix in ['.myo', '.mdb']:
            return self._extract_myob_content(file_path)
        elif suffix == '.txt':
            return self._extract_text_content(file_path)
        elif suffix.lower() in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff']:
            return self._extract_image_content(file_path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

    def _extract_pdf_content(self, file_path: Path) -> str:
        """Extract text from PDF files"""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text()
                
                self.extraction_stats['pdf']['success'] += 1
                return text
        except Exception as e:
            self.extraction_stats['pdf']['failed'] += 1
            raise Exception(f"PDF extraction failed: {str(e)}")

    def _extract_doc_content(self, file_path: Path) -> str:
        """Extract text from DOC/DOCX files"""
        try:
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            self.extraction_stats['doc']['success'] += 1
            return text
        except Exception as e:
            self.extraction_stats['doc']['failed'] += 1
            raise Exception(f"DOC extraction failed: {str(e)}")

    def _extract_excel_content(self, file_path: Path) -> str:
        """Extract text from modern Excel files (.xlsx)"""
        try:
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            text = []
            
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                for row in ws.rows:
                    row_text = [str(cell.value) if cell.value is not None else "" for cell in row]
                    text.append(" ".join(row_text))
            
            self.extraction_stats['xlsx']['success'] += 1
            return "\n".join(text)
        except Exception as e:
            self.extraction_stats['xlsx']['failed'] += 1
            raise Exception(f"Excel extraction failed: {str(e)}")

    def _extract_legacy_excel_content(self, file_path: Path) -> str:
        """Extract text from legacy Excel files (.xls)"""
        try:
            wb = xlrd.open_workbook(str(file_path))
            text = []
            
            for sheet_idx in range(wb.nsheets):
                sheet = wb.sheet_by_index(sheet_idx)
                for row_idx in range(sheet.nrows):
                    row_text = [str(cell.value) if cell.value else "" for cell in sheet.row(row_idx)]
                    text.append(" ".join(row_text))
            
            self.extraction_stats['xls']['success'] += 1
            return "\n".join(text)
        except Exception as e:
            self.extraction_stats['xls']['failed'] += 1
            raise Exception(f"Legacy Excel extraction failed: {str(e)}")

    def _extract_csv_content(self, file_path: Path) -> str:
        """Extract text from CSV files"""
        try:
            # Use pandas to handle various CSV formats and encodings
            df = pd.read_csv(file_path)
            text = df.to_string(index=False)
            
            self.extraction_stats['csv']['success'] += 1
            return text
        except Exception as e:
            self.extraction_stats['csv']['failed'] += 1
            raise Exception(f"CSV extraction failed: {str(e)}")

    def _detect_faces(self, image_path: str, image: np.ndarray) -> List[Dict]:
        """Detect and analyze faces in the image"""
        faces_info = []
        try:
            # Convert BGR to RGB (face_recognition uses RGB)
            rgb_image = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
            
            # Find face locations and encodings
            face_locations = face_recognition.face_locations(rgb_image)
            face_encodings = face_recognition.face_encodings(rgb_image, face_locations)
            
            # Analyze each face
            for idx, (face_loc, face_enc) in enumerate(zip(face_locations, face_encodings)):
                top, right, bottom, left = face_loc
                face_image = rgb_image[top:bottom, left:right]
                
                # Convert face image for DeepFace
                face_image_bgr = cv2.cvtColor(face_image, cv2.COLOR_RGB2BGR)
                
                # Analyze demographics and emotions
                analysis = DeepFace.analyze(face_image_bgr, 
                                          actions=['age', 'gender', 'race', 'emotion'],
                                          enforce_detection=False)
                
                # Create face signature for potential matching
                face_signature = hashlib.sha256(face_enc.tobytes()).hexdigest()
                
                face_info = {
                    'location': {
                        'top': top,
                        'right': right,
                        'bottom': bottom,
                        'left': left
                    },
                    'analysis': {
                        'age': analysis[0]['age'],
                        'gender': analysis[0]['gender'],
                        'race': analysis[0]['dominant_race'],
                        'emotion': analysis[0]['dominant_emotion']
                    },
                    'signature': face_signature
                }
                faces_info.append(face_info)
                
        except Exception as e:
            print(f"Face detection error for {image_path}: {str(e)}")
        
        return faces_info

    def _extract_metadata(self, file_path: Path) -> Dict:
        """Extract comprehensive metadata from file"""
        stats = file_path.stat()
        metadata = {
            'basic': {
                'filename': file_path.name,
                'size': stats.st_size,
                'created': datetime.fromtimestamp(stats.st_ctime).isoformat(),
                'modified': datetime.fromtimestamp(stats.st_mtime).isoformat(),
                'type': file_path.suffix.lower()
            },
            'exif': {},
            'mime': None,
            'source': None
        }
        
        try:
            # Get MIME type
            mime = magic.Magic(mime=True)
            metadata['mime'] = mime.from_file(str(file_path))
            
            # Extract EXIF data for images
            if metadata['mime'].startswith('image/'):
                with open(file_path, 'rb') as f:
                    tags = exifread.process_file(f, details=False)
                    
                    # Convert EXIF tags to serializable format
                    exif_data = {}
                    for tag, value in tags.items():
                        if tag not in ('JPEGThumbnail', 'TIFFThumbnail', 'Filename', 'EXIF MakerNote'):
                            exif_data[tag] = str(value)
                    
                    metadata['exif'] = exif_data
            
            # Detect source application/platform
            metadata['source'] = self._detect_source(file_path, metadata)
            
        except Exception as e:
            print(f"Metadata extraction error for {file_path}: {str(e)}")
        
        return metadata

    def _detect_source(self, file_path: Path, metadata: Dict) -> Optional[str]:
        """Detect the source application or platform of the file"""
        # Check file content
        try:
            with open(file_path, 'rb') as f:
                content = f.read(4096)  # Read first 4KB
                content_str = content.decode('utf-8', errors='ignore').lower()
                
                # WhatsApp indicators
                if any(indicator in content_str for indicator in [
                    'whatsapp', 'wa.me', 'message from whatsapp',
                    '[\d{1,2}:\d{2}:\d{2} [AP]M]'  # WhatsApp timestamp pattern
                ]):
                    return 'whatsapp'
                
                # Instagram indicators
                if any(indicator in content_str for indicator in [
                    'instagram.com', 'ig direct', '@instagram',
                    'instagram story', 'instagram post'
                ]):
                    return 'instagram'
                
                # Facebook indicators
                if any(indicator in content_str for indicator in [
                    'facebook.com', 'fb.com', 'messenger.com',
                    'facebook messenger'
                ]):
                    return 'facebook'
                
                # Twitter/X indicators
                if any(indicator in content_str for indicator in [
                    'twitter.com', 'x.com', '@twitter',
                    'tweet', 'twitter web app'
                ]):
                    return 'twitter'
                
                # LinkedIn indicators
                if any(indicator in content_str for indicator in [
                    'linkedin.com', 'lnkd.in',
                    'linkedin post', 'linkedin message'
                ]):
                    return 'linkedin'
                
                # Email indicators
                if any(indicator in content_str for indicator in [
                    'from:', 'to:', 'subject:', 'date:',
                    'content-type: text/html',
                    'message-id:', 'mime-version:'
                ]):
                    return 'email'
                
        except Exception:
            pass
        
        # Check EXIF metadata
        if metadata.get('exif'):
            exif_str = json.dumps(metadata['exif']).lower()
            
            # Check software/creator tags
            if 'whatsapp' in exif_str:
                return 'whatsapp'
            if 'instagram' in exif_str:
                return 'instagram'
            if 'facebook' in exif_str:
                return 'facebook'
            if 'twitter' in exif_str:
                return 'twitter'
            if 'linkedin' in exif_str:
                return 'linkedin'
            if 'outlook' in exif_str or 'thunderbird' in exif_str:
                return 'email'
        
        return None

    def _extract_dates(self, content: str) -> List[str]:
        """Extract and normalize dates from content"""
        date_patterns = [
            # Standard formats
            r'\d{2}[-/]\d{2}[-/]\d{4}',  # DD-MM-YYYY
            r'\d{4}[-/]\d{2}[-/]\d{2}',  # YYYY-MM-DD
            r'\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}',  # 1 Jan 2020
            
            # Social media timestamps
            r'\d{1,2}:\d{2}(?::\d{2})?\s*(?:AM|PM)?\s*[-–]\s*\d{1,2}/\d{1,2}/\d{2,4}',  # 12:34 PM - 01/02/2020
            r'\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+(?:at\s+)?\d{1,2}:\d{2}(?::\d{2})?\s*(?:AM|PM)?',  # 1 Jan at 12:34 PM
            
            # Message/email timestamps
            r'(?:Sent|Received|Posted|Created):\s*\d{1,2}[-/]\d{1,2}[-/]\d{2,4}\s*(?:at\s+)?\d{1,2}:\d{2}(?::\d{2})?\s*(?:AM|PM)?',
            
            # Relative timestamps (with conversion logic)
            r'(?:Today|Yesterday)\s+(?:at\s+)?\d{1,2}:\d{2}(?::\d{2})?\s*(?:AM|PM)?',
            r'\d+\s+(?:minute|hour|day|week|month|year)s?\s+ago'
        ]
        
        dates = []
        for pattern in date_patterns:
            matches = re.finditer(pattern, content)
            dates.extend(match.group(0) for match in matches)
        
        # Convert relative dates to absolute dates
        now = datetime.now()
        normalized_dates = []
        for date_str in dates:
            try:
                if 'ago' in date_str.lower():
                    # Convert relative time to absolute
                    amount, unit, _ = date_str.lower().split()
                    amount = int(amount)
                    if 'minute' in unit:
                        date = now - timedelta(minutes=amount)
                    elif 'hour' in unit:
                        date = now - timedelta(hours=amount)
                    elif 'day' in unit:
                        date = now - timedelta(days=amount)
                    elif 'week' in unit:
                        date = now - timedelta(weeks=amount)
                    elif 'month' in unit:
                        date = now - timedelta(days=amount*30)  # Approximate
                    elif 'year' in unit:
                        date = now - timedelta(days=amount*365)  # Approximate
                    normalized_dates.append(date.strftime('%Y-%m-%d %H:%M:%S'))
                elif 'today' in date_str.lower():
                    time_part = re.search(r'\d{1,2}:\d{2}(?::\d{2})?\s*(?:AM|PM)?', date_str)
                    if time_part:
                        normalized_dates.append(f"{now.strftime('%Y-%m-%d')} {time_part.group(0)}")
                elif 'yesterday' in date_str.lower():
                    yesterday = now - timedelta(days=1)
                    time_part = re.search(r'\d{1,2}:\d{2}(?::\d{2})?\s*(?:AM|PM)?', date_str)
                    if time_part:
                        normalized_dates.append(f"{yesterday.strftime('%Y-%m-%d')} {time_part.group(0)}")
                else:
                    normalized_dates.append(date_str)
            except Exception:
                normalized_dates.append(date_str)
        
        return sorted(list(set(normalized_dates)))

    def _preprocess_image(self, image: np.ndarray) -> np.ndarray:
        """Preprocess image for better OCR results"""
        # Convert to grayscale
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        
        # Apply adaptive thresholding
        binary = cv2.adaptiveThreshold(
            gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
        )
        
        # Noise removal
        denoised = cv2.fastNlMeansDenoising(binary)
        
        # Deskew if needed
        coords = np.column_stack(np.where(denoised > 0))
        angle = cv2.minAreaRect(coords)[-1]
        if angle < -45:
            angle = 90 + angle
        center = tuple(np.array(denoised.shape[1::-1]) / 2)
        rot_mat = cv2.getRotationMatrix2D(center, angle, 1.0)
        result = cv2.warpAffine(denoised, rot_mat, denoised.shape[1::-1],
                               flags=cv2.INTER_LINEAR, borderMode=cv2.BORDER_REPLICATE)
        
        return result

    def _detect_image_type(self, image: np.ndarray, text: str) -> str:
        """Detect the type of image based on content and visual features"""
        # Convert text to lowercase for easier matching
        text_lower = text.lower()
        
        # Gift card indicators
        gift_card_keywords = ['gift card', 'gift certificate', 'balance', 'redeem',
                            'value', 'card number', 'pin', 'activation']
        
        # Transaction screenshot indicators
        transaction_keywords = ['transaction', 'payment', 'transfer', 'amount',
                              'reference', 'account', 'date', 'balance',
                              'debit', 'credit', 'receipt']
        
        # Document indicators
        document_keywords = ['confidential', 'agreement', 'contract', 'report',
                           'policy', 'letter', 'statement', 'invoice']
        
        # Count matches for each type
        gift_card_score = sum(1 for kw in gift_card_keywords if kw in text_lower)
        transaction_score = sum(1 for kw in transaction_keywords if kw in text_lower)
        document_score = sum(1 for kw in document_keywords if kw in text_lower)
        
        # Additional visual analysis for gift cards
        # Gift cards often have more uniform color distribution
        color_hist = cv2.calcHist([image], [0], None, [256], [0, 256])
        color_variance = np.var(color_hist)
        if color_variance < 1000:  # Threshold determined empirically
            gift_card_score += 2
        
        # Determine the type based on scores
        max_score = max(gift_card_score, transaction_score, document_score)
        if max_score == 0:
            return 'unknown'
        elif max_score == gift_card_score:
            return 'gift_card'
        elif max_score == transaction_score:
            return 'transaction'
        else:
            return 'document'

    def _extract_image_content(self, file_path: Path) -> str:
        """Extract text from images using OCR with specialized processing"""
        try:
            # Read image using OpenCV
            image = cv2.imread(str(file_path))
            if image is None:
                raise ValueError(f"Could not read image: {file_path}")
            
            # Detect faces
            faces = self._detect_faces(str(file_path), image)
            
            # Preprocess image
            processed_image = self._preprocess_image(image)
            
            # Perform OCR
            text = pytesseract.image_to_string(processed_image)
            
            # Extract metadata including EXIF
            metadata = self._extract_metadata(file_path)
            
            # Detect image type
            image_type = self._detect_image_type(image, text)
            
            # Extract dates from text
            dates = self._extract_dates(text)
            
            # Build result dictionary
            result = {
                'type': image_type,
                'faces': faces,
                'dates': dates,
                'metadata': metadata,
                'source': metadata['source'],
                'content': text
            }
            
            # Additional processing based on image type
            if image_type == 'gift_card':
                # Extract card number and balance
                card_numbers = re.findall(r'\b\d{4}[- ]?\d{4}[- ]?\d{4}[- ]?\d{4}\b', text)
                balance = re.findall(r'\$\s*\d+(?:\.\d{2})?', text)
                
                result = [f"Type: Gift Card"]
                if card_numbers:
                    result.append(f"Card Number: {card_numbers[0]}")
                if balance:
                    result.append(f"Balance: {balance[0]}")
                result.append("\nRaw Text:")
                result.append(text)
                
                text = "\n".join(result)
                
            elif image_type == 'transaction':
                # Extract transaction details
                dates = re.findall(r'\d{1,2}[-/]\d{1,2}[-/]\d{2,4}', text)
                amounts = re.findall(r'\$\s*\d+(?:\.\d{2})?', text)
                references = re.findall(r'ref[:\s]*(\w+)', text, re.IGNORECASE)
                
                result = [f"Type: Transaction Screenshot"]
                if dates:
                    result.append(f"Date: {dates[0]}")
                if amounts:
                    result.append(f"Amount: {amounts[0]}")
                if references:
                    result.append(f"Reference: {references[0]}")
                result.append("\nRaw Text:")
                result.append(text)
                
                text = "\n".join(result)
                
            else:  # document or unknown
                text = f"Type: {'Document' if image_type == 'document' else 'Unknown'}\n\n" + text
            
            self.extraction_stats['img']['success'] += 1
            return text
            
        except Exception as e:
            self.extraction_stats['img']['failed'] += 1
            raise Exception(f"Image extraction failed: {str(e)}")

    def _extract_myob_content(self, file_path: Path) -> str:
        """Extract text from MYOB database files"""
        try:
            # Connect to MYOB database using ODBC
            conn_str = f"DRIVER={{Microsoft Access Driver (*.mdb, *.accdb)}};DBQ={file_path};"
            conn = pyodbc.connect(conn_str)
            cursor = conn.cursor()
            
            # Get list of tables
            tables = [row.table_name for row in cursor.tables(tableType='TABLE')]
            
            # Extract data from each table
            text = []
            for table in tables:
                df = pd.read_sql(f"SELECT * FROM [{table}]", conn)
                text.append(f"Table: {table}")
                text.append(df.to_string(index=False))
                text.append("\n")
            
            conn.close()
            self.extraction_stats['myo']['success'] += 1
            return "\n".join(text)
        except Exception as e:
            self.extraction_stats['myo']['failed'] += 1
            raise Exception(f"MYOB extraction failed: {str(e)}")

    def _extract_text_content(self, file_path: Path) -> str:
        """Extract text from plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            self.extraction_stats['txt']['success'] += 1
            return text
        except Exception as e:
            self.extraction_stats['txt']['failed'] += 1
            raise Exception(f"Text extraction failed: {str(e)}")

    def _extract_metadata(self, file_path: Path) -> Dict:
        """Extract metadata from file"""
        stats = file_path.stat()
        return {
            'filename': file_path.name,
            'size': stats.st_size,
            'created': datetime.fromtimestamp(stats.st_ctime).isoformat(),
            'modified': datetime.fromtimestamp(stats.st_mtime).isoformat(),
            'type': file_path.suffix.lower()
        }

    def _extract_entities(self, content: str) -> Dict[str, list]:
        """Extract various entity types from content"""
        entities = {
            'companies': [],
            'people': [],
            'locations': [],
            'dates': []
        }
        
        # Company patterns
        company_patterns = [
            r'(?i)(?:Pty\.?\s*Ltd\.?|Limited|Inc\.?|Corporation)',
            r'(?i)(?:Trading\s+as|T/A)\s+([A-Z][A-Za-z\s]+)',
            r'(?i)(?:ACN|ABN|ARBN)\s*:\s*(\d{9}|\d{11})'
        ]
        
        # Person patterns
        person_patterns = [
            r'(?i)(?:Mr\.|Mrs\.|Ms\.|Dr\.)\s+([A-Z][a-z]+\s+[A-Z][a-z]+)',
            r'(?i)(?:Director|Secretary|Trustee):\s+([A-Z][a-z]+\s+[A-Z][a-z]+)'
        ]
        
        # Process patterns
        for pattern in company_patterns:
            matches = re.finditer(pattern, content)
            entities['companies'].extend(match.group(0) for match in matches)
            
        for pattern in person_patterns:
            matches = re.finditer(pattern, content)
            entities['people'].extend(match.group(1) for match in matches if match.group(1))
        
        return entities

    def _extract_dates(self, content: str) -> list:
        """Extract dates from content"""
        date_patterns = [
            r'\d{2}[-/]\d{2}[-/]\d{4}',  # DD-MM-YYYY
            r'\d{4}[-/]\d{2}[-/]\d{2}',  # YYYY-MM-DD
            r'\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}'  # 1 Jan 2020
        ]
        
        dates = []
        for pattern in date_patterns:
            matches = re.finditer(pattern, content)
            dates.extend(match.group(0) for match in matches)
            
        return sorted(list(set(dates)))

    def generate_processing_report(self) -> Dict:
        """Generate a report of processing statistics"""
        return {
            'total_files': len(self.processed_files),
            'extraction_stats': self.extraction_stats,
            'file_types': {
                suffix: len([f for f in self.processed_files if f.endswith(suffix)])
                for suffix in ['.pdf', '.doc', '.docx', '.xls', '.xlsx', '.csv', '.txt', '.myo', '.mdb',
                               '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff']
            },
            'generated_at': datetime.now().isoformat()
        }